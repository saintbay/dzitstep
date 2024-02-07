
import pandas as pd
import openpyxl
import json
from docx import Document
from docx.shared import Pt
from openpyxl.styles import Font, Border, Side

# Задание 1а
df1 = pd.DataFrame({'data': [1111]})
df2 = pd.DataFrame({'data': [2222]})
df3 = pd.DataFrame({'data': [3333]})

# Задание 1б
merged_df = pd.concat([df1, df2, df3], ignore_index=True)
sorted_df = merged_df.sort_values(by='data', ascending=False)

# Задание 1в
with pd.ExcelWriter('result.xlsx', engine='openpyxl') as writer:
    sorted_df.to_excel(writer, index=False, sheet_name='Sheet1')
    workbook = writer.book
    worksheet = writer.sheets['Sheet1']

    font = Font(name='Arial', size=12, bold=True)
    border = Border(left=Side(border_style='thin'), right=Side(border_style='thin'), top=Side(border_style='thin'), bottom=Side(border_style='thin'))

    for cell in worksheet['1']:
        cell.font = font
        cell.border = border

# Задание 1а
with open('data.json', 'w') as json_file:
    json_file.write(json.dumps({'data': 'Hello Python'}))

# Задание 1б
with open('data.json', 'r') as json_file:
    data_array = json.load(json_file)

# Задание 1в
for i, data_dict in enumerate(data_array):
    with open(f'data_{i}.json', 'w') as json_file:
        json_file.write(json.dumps(data_dict))

# Задание 1а
document = Document()
document.add_paragraph('Hello Python', style='Heading1')

# Задание 1б
bold_text = ''
for paragraph in document.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            bold_text += run.text

# Задание 1в
new_document = Document()
new_paragraph = new_document.add_paragraph('New paragraph with different font and size.')
new_run = new_paragraph.runs[0]
new_run.font.name = 'Times New Roman'
new_run.font.size = Pt(14)
new_document.save('new_document.docx')
